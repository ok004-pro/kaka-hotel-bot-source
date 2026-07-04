# -*- coding: utf-8 -*-
"""
호텔 지식창고 통합 빌더 (v2.0)
역할: "카카" (Kaka)

기존 규약집 검색을 뛰어넘어,
- 집합건물법 법령해설 전문 (.docx)
- 위탁운영계약서 및 생활숙박시설 계약서 전문 (.docx)
- 집합건물 관리 업무 매뉴얼 및 감사자료/분쟁사례 (.pdf)
을 모두 파싱하여 RAG 요약 지식 소스인 'hotel_regulation.txt'에 똑똑하게 인출/이식합니다.
"""

import json
from pathlib import Path
import docx
from pypdf import PdfReader

BASE_DIR = Path("C:/Users/PC/hotel_bot")
KNOWLEDGE_DIR = BASE_DIR / "knowledge"
OUTPUT_TXT = KNOWLEDGE_DIR / "hotel_regulation.txt"

def extract_docx(path):
    doc = docx.Document(path)
    text = []
    text.append(f"\n==================================================")
    text.append(f"📄문서 출처: {path.name}")
    text.append(f"==================================================\n")
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                text.append(" | ".join(row_text))
    return "\n".join(text)

def extract_pdf(path):
    reader = PdfReader(path)
    text = []
    text.append(f"\n==================================================")
    text.append(f"📄문서 출처: {path.name}")
    text.append(f"==================================================\n")
    # 대용량 PDF의 경우, 요약 및 RAG 매칭을 위해 핵심 부분 추출
    # 대량 데이터가 들어가므로 상위 50페이지 등으로 제어하거나 전문 추출
    max_pages = 80 if "매뉴얼" in path.name or "가이드" in path.name else 200
    pages_to_read = min(len(reader.pages), max_pages)
    
    print(f"PDF {path.name} (총 {len(reader.pages)}쪽 중 {pages_to_read}쪽 추출)")
    for i in range(pages_to_read):
        page_text = reader.pages[i].extract_text()
        if page_text and page_text.strip():
            text.append(f"--- [파일: {path.name} / {i+1} Page] ---")
            text.append(page_text.strip())
    return "\n".join(text)

def main():
    files = list(KNOWLEDGE_DIR.glob("*"))
    combined_text = []

    # 1. 기존 가상 규약 텍스트 및 기본 규약 먼저 탑재
    print("통합 지식 이식 가동 시작...")
    
    # 2. 업로드해주신 docx 및 pdf 우선 인출
    for f in files:
        if f.suffix.lower() == ".docx":
            try:
                txt = extract_docx(f)
                combined_text.append(txt)
                print(f"✅ {f.name} 완벽 추출 성공 (크기: {len(txt)} 자)")
            except Exception as e:
                print(f"❌ {f.name} 오류: {e}")
        elif f.suffix.lower() == ".pdf":
            # 무거운 PDF 중 대자료도 안전하게 파싱
            try:
                txt = extract_pdf(f)
                combined_text.append(txt)
                print(f"✅ {f.name} 완벽 추출 성공 (크기: {len(txt)} 자)")
            except Exception as e:
                print(f"❌ {f.name} 오류: {e}")

    if combined_text:
        full_source = "\n\n".join(combined_text)
        OUTPUT_TXT.write_text(full_source, encoding="utf-8")
        print(f"\n🏆 통합 대용량 지식창고 이식 빌드 대성공!")
        print(f"저장 경로: {OUTPUT_TXT}")
        print(f"총 분석 데이터량: {len(full_source)} 자")
    else:
        print("이식할 문서를 발견하지 못했습니다.")

if __name__ == "__main__":
    main()
